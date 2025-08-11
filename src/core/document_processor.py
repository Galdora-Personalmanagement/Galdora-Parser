import io
import os
from PIL import Image
import pytesseract
import PyPDF2
from pdf2image import convert_from_path
from docx import Document

class DocumentProcessor:
    """Klasse zur Verarbeitung verschiedener Dokumenttypen"""
    
    def process_document(self, file_path, file_extension):
        """
        Extrahiert Text aus Dokumenten verschiedenen Typs
        
        Args:
            file_path: Pfad zur Datei
            file_extension: Dateierweiterung
        
        Returns:
            String mit extrahiertem Text
        """
        if file_extension.lower() in ['.pdf']:
            return self._extract_from_pdf(file_path)
        elif file_extension.lower() in ['.jpg', '.jpeg', '.png']:
            return self._extract_from_image(file_path)
        elif file_extension.lower() in ['.docx']:
            return self._extract_from_docx(file_path)
        else:
            raise ValueError(f"Nicht unterstützter Dateityp: {file_extension}")
    
    def _extract_from_pdf(self, file_path):
        """Extrahiert Text aus PDF-Dateien"""
        text = ""
        
        # Versuche zunächst direkte Textextraktion
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:  # Wenn Text erfolgreich extrahiert wurde
                        text += page_text + "\n"
        except Exception as e:
            print(f"Fehler bei direkter PDF-Textextraktion: {str(e)}")
        
        # Falls kein oder wenig Text extrahiert wurde, OCR verwenden (speicherschonend, seitenweise)
        if len(text.strip()) < 100:  # Heuristik: Weniger als 100 Zeichen bedeutet wahrscheinlich ein Scan
            try:
                from pdf2image import pdfinfo_from_path
                import tempfile
                info = pdfinfo_from_path(file_path)
                num_pages = int(info.get('Pages', 0))
                if num_pages == 0:
                    # Fallback: Einmalige Konvertierung versuchen
                    images = convert_from_path(file_path, dpi=200, thread_count=1)
                    for image in images:
                        try:
                            page_text = pytesseract.image_to_string(image, lang='deu')
                            text += page_text + "\n"
                        finally:
                            try:
                                image.close()
                            except Exception:
                                pass
                else:
                    # Seitenweise konvertieren, um RAM zu sparen
                    with tempfile.TemporaryDirectory() as tmpdir:
                        for page in range(1, num_pages + 1):
                            try:
                                paths = convert_from_path(
                                    file_path,
                                    dpi=200,
                                    first_page=page,
                                    last_page=page,
                                    output_folder=tmpdir,
                                    fmt='png',
                                    thread_count=1,
                                    paths_only=True
                                )
                                if not paths:
                                    continue
                                img_path = paths[0]
                                try:
                                    from PIL import Image as PILImage
                                    img = PILImage.open(img_path)
                                    page_text = pytesseract.image_to_string(img, lang='deu')
                                    text += page_text + "\n"
                                finally:
                                    try:
                                        img.close()
                                    except Exception:
                                        pass
                                    try:
                                        os.remove(img_path)
                                    except Exception:
                                        pass
                            except Exception as per_page_err:
                                print(f"Fehler bei PDF-OCR (Seite {page}): {per_page_err}")
                                continue
            except Exception as e:
                print(f"Fehler bei PDF-OCR: {str(e)}")
                # Wenn auch OCR fehlschlägt, zurückgeben was wir haben
        
        return text
    
    def _extract_from_image(self, file_path):
        """Extrahiert Text aus Bilddateien mit OCR"""
        try:
            image = Image.open(file_path)
            # OCR für deutsche Dokumente (für andere Sprachen anpassen)
            text = pytesseract.image_to_string(image, lang='deu')
            return text
        except Exception as e:
            raise Exception(f"Fehler bei der Bildverarbeitung: {str(e)}")
        finally:
            try:
                image.close()
            except Exception:
                pass
    
    def _extract_from_docx(self, file_path):
        """Extrahiert Text aus Word-Dokumenten"""
        try:
            doc = Document(file_path)
            full_text = []
            
            # Text aus Absätzen extrahieren
            for para in doc.paragraphs:
                full_text.append(para.text)
            
            # Text aus Tabellen extrahieren
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text.append(cell.text)
            
            return '\n'.join(full_text)
        except Exception as e:
            raise Exception(f"Fehler bei der Word-Dokumentverarbeitung: {str(e)}")

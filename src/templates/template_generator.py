from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, HRFlowable, PageBreak, KeepTogether, Frame, PageTemplate, NextPageTemplate
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm, cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import io
import os
import docx
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from src.utils.image_utils import get_image_path, ensure_images_in_static
from src.utils.company_config import get_company_config, get_company_logo_path
import json
from PIL import Image as PILImage

class ProfileGenerator:
    """Klasse zur Erstellung von PDF-Profilen im Classic Template Design"""
    
    def __init__(self, selected_company="galdora"):
        """Initialisiert den Profil-Generator mit Stilen und Unternehmen-Konfiguration"""
        self.styles = getSampleStyleSheet()
        self.custom_styles = self._create_custom_styles()
        self.selected_company = selected_company
        self.company_config = get_company_config(selected_company)
    
    def generate_profile(self, profile_data, output_path, template="classic", format="pdf"):
        """
        Generiert ein Profil aus den extrahierten Daten im Classic Template Format
        
        Args:
            profile_data: Dictionary mit Profildaten
            output_path: Pfad für die Ausgabedatei
            template: Art der Vorlage (nur "classic" verfügbar)
            format: Format der Ausgabedatei ("pdf" oder "docx")
        
        Returns:
            Pfad zur generierten Datei
        """
        if not isinstance(profile_data, dict):
            print(f"Warnung: profile_data ist kein Dictionary. Typ: {type(profile_data)}")
            profile_data = {}
        
        if not output_path:
            raise ValueError("Der Ausgabepfad darf nicht leer sein.")
            
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
        
        if format.lower() == "pdf":
            return self._generate_pdf_profile(profile_data, output_path)
        elif format.lower() == "docx":
            if not output_path.lower().endswith('.docx'):
                base_path = os.path.splitext(output_path)[0]
                output_path = base_path + '.docx'
            return self._generate_docx_profile(profile_data, output_path)
        else:
            raise ValueError(f"Ungültiges Format: {format}. Unterstützte Formate: 'pdf', 'docx'")
    
    def _generate_pdf_profile(self, profile_data, output_path):
        """Generiert ein PDF-Profil aus den extrahierten Daten im Classic Template Design"""
        try:
            doc = SimpleDocTemplate(
                output_path,
                pagesize=A4,
                rightMargin=20*mm, 
                leftMargin=20*mm,
                topMargin=20*mm, 
                bottomMargin=25*mm
            )
            
            elements = self._create_document_elements(profile_data)
            
            def footer_callback(canvas, doc):
                canvas.saveState()
                footer_style = self.custom_styles['Footer']
                
                y_position = 15*mm
                
                canvas.setStrokeColor(colors.lightgrey)
                canvas.setLineWidth(1)
                canvas.line(20*mm, y_position + 5*mm, A4[0] - 20*mm, y_position + 5*mm)
                
                footer_text = self._get_dynamic_footer_text()
                p = Paragraph(footer_text, footer_style)
                w, h = p.wrap(A4[0] - 40*mm, 20*mm)
                p.drawOn(canvas, 20*mm, y_position - h + 2*mm)
                
                page_num = canvas.getPageNumber()
                canvas.setFont("Helvetica", 8)
                canvas.setFillColor(colors.grey)
                canvas.drawRightString(A4[0] - 20*mm, y_position - h - 3*mm, f"Seite {page_num}")
                
                canvas.restoreState()
            
            first_page_frame = Frame(
                doc.leftMargin, 
                doc.bottomMargin, 
                A4[0] - doc.leftMargin - doc.rightMargin, 
                A4[1] - doc.topMargin - doc.bottomMargin
            )
            
            first_page_template = PageTemplate(
                id='firstPage', 
                frames=[first_page_frame], 
                onPage=footer_callback
            )
            
            later_page_frame = Frame(
                doc.leftMargin, 
                doc.bottomMargin, 
                A4[0] - doc.leftMargin - doc.rightMargin, 
                A4[1] - doc.topMargin - doc.bottomMargin
            )
            
            later_pages_template = PageTemplate(
                id='laterPages', 
                frames=[later_page_frame], 
                onPage=footer_callback
            )
            
            doc.addPageTemplates([first_page_template, later_pages_template])
            
            elements.insert(0, NextPageTemplate('laterPages'))
            
            doc.build(elements)
            
            return output_path
        except Exception as e:
            print(f"Fehler bei der PDF-Generierung: {str(e)}")
            raise

    def _generate_docx_profile(self, profile_data, output_path):
        """Generiert ein Word-Dokument (DOCX) aus den extrahierten Daten im Classic Template Design"""
        try:
            doc = docx.Document()
            
            section = doc.sections[0]
            section.page_width = Cm(21)
            section.page_height = Cm(29.7)
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            
            logo_style = doc.styles.add_style('GaldoraLogo', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
            logo_style.font.size = Pt(36)
            logo_style.font.bold = True
            
            ensure_images_in_static()
            
            logo_path = self._get_company_logo_path()
            
            if logo_path and os.path.exists(logo_path):
                try:
                    paragraph = doc.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run = paragraph.add_run()
                    run.add_picture(logo_path, width=Cm(5))
                except Exception as e:
                    print(f"DOCX: Fehler beim Laden des Logos: {str(e)}")
                    company_name = self.company_config.get('name', 'GALDORA')
                    doc.add_paragraph(company_name, style='GaldoraLogo').alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                print(f"DOCX: Logo-Datei nicht gefunden: {logo_path}")
                company_name = self.company_config.get('name', 'GALDORA')
                doc.add_paragraph(company_name, style='GaldoraLogo').alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            doc.add_paragraph()
            
            title_style = doc.styles.add_style('ProfilTitle', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.size = Pt(16)
            title_style.font.bold = True
            title_style.font.color.rgb = RGBColor(25, 115, 184)
            
            doc.add_paragraph("Profil", style='ProfilTitle')
            
            self._create_classic_docx_content(doc, profile_data)
            
            doc.save(output_path)
            print(f"DOCX-Datei erfolgreich erstellt: {output_path}")
            
            return output_path
            
        except Exception as e:
            print(f"Fehler bei der DOCX-Generierung: {str(e)}")
            raise e

    def _create_document_elements(self, profile_data):
        """Erstellt die Dokumentelemente für das Classic Template"""
        elements = []
        try:
            ensure_images_in_static()
            
            logo_path = self._get_company_logo_path()

            if logo_path and os.path.exists(logo_path):
                try:
                    img_pil = PILImage.open(logo_path)
                    img_width, img_height = img_pil.size
                    aspect_ratio = img_width / img_height
                    
                    if self.selected_company == "bejob":
                        target_width = 180
                        target_height = target_width / aspect_ratio
                        img = Image(logo_path, width=target_width, height=target_height)
                        logo_table = Table([["", img]], colWidths=[A4[0] - target_width - 50*mm, target_width + 10])
                        logo_table.setStyle(TableStyle([
                            ('ALIGN', (1, 0), (1, 0), 'RIGHT'),
                            ('VALIGN', (0, 0), (1, 0), 'TOP'),
                        ]))
                        elements.append(logo_table)
                        elements.append(Spacer(1, 0.1*cm))
                    else:
                        target_width = 160
                        target_height = target_width / aspect_ratio
                        img = Image(logo_path, width=target_width, height=target_height)
                        logo_table = Table([[img, ""]], colWidths=[target_width + 10, A4[0] - target_width - 50*mm])
                        logo_table.setStyle(TableStyle([
                            ('ALIGN', (0, 0), (0, 0), 'LEFT'),
                            ('VALIGN', (0, 0), (1, 0), 'TOP'),
                        ]))
                        elements.append(logo_table)
                        elements.append(Spacer(1, 0.2*cm))
                except Exception as e:
                    print(f"Fehler beim Laden des Logos: {str(e)}")
                    elements.append(Paragraph("GALDORA", self.custom_styles['GaldoraLogo']))
            else:
                print(f"Logo-Datei nicht gefunden oder ungültig: {logo_path}")
                elements.append(Paragraph("GALDORA", self.custom_styles['GaldoraLogo']))
            
            self._create_classic_content(elements, profile_data)
            
            return elements
            
        except Exception as e:
            print(f"Fehler beim Erstellen der Dokumentelemente: {str(e)}")
            raise

    def _create_classic_content(self, elements, profile_data):
        """Erstellt den Content für das Classic Template"""
        
        personal_data = profile_data.get('persönliche_daten', {})
        
        if self.selected_company == "bejob":
            elements.append(Spacer(1, -0.2*cm))
        
        elements.append(Paragraph("Profil", self.custom_styles['ProfilTitle']))
        
        name = personal_data.get('name', 'Profil')
        elements.append(Paragraph(name, self.custom_styles['Name']))
        
        grayed_style = ParagraphStyle(
            'GrayedLabelInline',
            parent=self.custom_styles['LabelInline'],
            textColor=colors.grey
        )
        
        if personal_data.get('wohnort'):
            elements.append(Paragraph(f"Wohnort: {personal_data.get('wohnort')}", grayed_style))

        verfuegbarkeit_status = profile_data.get('verfuegbarkeit_status', '')
        verfuegbarkeit_details = profile_data.get('verfuegbarkeit_details', '')
        if verfuegbarkeit_status:
            if verfuegbarkeit_details:
                elements.append(Paragraph(f"ab {verfuegbarkeit_details}", grayed_style))
            else:
                elements.append(Paragraph(f"ab {verfuegbarkeit_status}", grayed_style))

        if personal_data.get('jahrgang'):
            elements.append(Paragraph(f"Jahrgang: {personal_data.get('jahrgang')}", grayed_style))

        fuehrerschein = personal_data.get('führerschein', '')
        if fuehrerschein and fuehrerschein.strip():
            elements.append(Paragraph(f"Führerschein: {fuehrerschein}", grayed_style))

        wunschgehalt = profile_data.get('wunschgehalt', '')
        if wunschgehalt:
            elements.append(Paragraph(f"Gehaltsvorstellung: {wunschgehalt}", grayed_style))

        elements.append(Spacer(1, 12))

        kontakt = personal_data.get('kontakt', {})
        ansprechpartner = kontakt.get('ansprechpartner', '')

        if ansprechpartner and ansprechpartner.strip() and ansprechpartner != "Kein Ansprechpartner":
            elements.append(HRFlowable(width="100%", thickness=0.5, color=colors.lightgrey, spaceBefore=3, spaceAfter=12))
            elements.append(Paragraph("IHR ANSPRECHPARTNER", self.custom_styles['ContactHeader']))
            
            if ansprechpartner == "Melike Demirkol":
                anrede = f"Frau Demirkol"
            elif ansprechpartner == "Alessandro Böhm":
                anrede = f"Herr Böhm"
                email = kontakt.get('email', "boehm@galdora.de")
            else:
                nachname = ansprechpartner.split()[-1] if ansprechpartner else 'Fischer'
                anrede = f"Herr {nachname}"
            
            if ansprechpartner == "Salim Alizai":
                telefon = kontakt.get('telefon', '')
            else:
                telefon = "02161 62126-00"
            
            if ansprechpartner == "Alessandro Böhm":
                email = kontakt.get("email", "boehm@galdora.de")
            elif ansprechpartner == "Salim Alizai":
                email = kontakt.get("email", "gl@galdora.de")
            elif ansprechpartner == "Konrad Ruszczyk":
                email = kontakt.get("email", "konrad@galdora.de")
            else:
                if 'nachname' in locals():
                    email = kontakt.get("email", f"{nachname.lower()}@galdora.de")
            
            grayed_contact_style = ParagraphStyle(
                'GrayedContactData',
                parent=self.custom_styles['ContactData'],
                textColor=colors.grey
            )
            
            elements.append(Paragraph(ansprechpartner, grayed_contact_style))
            
            if telefon:
                elements.append(Paragraph(f"Tel.: {telefon}", grayed_contact_style))
            
            if email:
                elements.append(Paragraph(f"E-Mail: {email}", grayed_contact_style))
            
            elements.append(Spacer(1, 12))
            
            elements.append(HRFlowable(width="100%", thickness=0.5, color=colors.lightgrey, spaceBefore=3, spaceAfter=12))
        
        berufserfahrung = profile_data.get('berufserfahrung', [])
        if berufserfahrung:
            elements.append(Paragraph("Beruflicher Werdegang", self.custom_styles['Heading2']))
            
            berufserfahrung_gruppen = []
            aktuelle_gruppe = []
            aktueller_platz = A4[1] - 400
            
            for i, erfahrung in enumerate(berufserfahrung):
                aufgaben_count = len(erfahrung.get('aufgaben', []))
                text_laenge = sum(len(aufgabe) for aufgabe in erfahrung.get('aufgaben', []))
                position_laenge = len(erfahrung.get('position', ''))
                unternehmen_laenge = len(erfahrung.get('unternehmen', ''))
                
                basis_hoehe = 40
                aufgaben_hoehe = min(20 * aufgaben_count, 10 + (text_laenge / 50))
                zusatz_hoehe = (position_laenge + unternehmen_laenge) / 100
                
                entry_height = basis_hoehe + aufgaben_hoehe + zusatz_hoehe
                
                if entry_height <= aktueller_platz:
                    aktuelle_gruppe.append((i, erfahrung, entry_height))
                    aktueller_platz -= entry_height
                else:
                    if aktuelle_gruppe:
                        berufserfahrung_gruppen.append(aktuelle_gruppe)
                    
                    aktuelle_gruppe = [(i, erfahrung, entry_height)]
                    aktueller_platz = A4[1] - 100 - entry_height
            
            if aktuelle_gruppe:
                berufserfahrung_gruppen.append(aktuelle_gruppe)
            
            for gruppen_index, gruppe in enumerate(berufserfahrung_gruppen):
                if gruppen_index > 0:
                    elements.append(PageBreak())
                
                gruppen_elemente = []
                
                for i, erfahrung, _ in gruppe:
                    zeitraum = erfahrung.get('zeitraum', '')
                    if "bis jetzt" in zeitraum.lower() or "bis heute" in zeitraum.lower():
                        zeitraum = zeitraum.replace("bis jetzt", "2025").replace("bis heute", "2025").replace("bis JETZT", "2025")
                    
                    unternehmen = erfahrung.get('unternehmen', '')
                    position = erfahrung.get('position', '')
                    
                    aufgaben = erfahrung.get('aufgaben', [])
                    
                    max_tasks = 6
                    if len(gruppe) > 2 and len(aufgaben) > 4:
                        aufgaben = aufgaben[:4]
                    elif len(aufgaben) > max_tasks:
                        aufgaben = aufgaben[:max_tasks]
                    
                    aufgaben_formatted = [Paragraph(f"• {aufgabe}", self.custom_styles['Normal']) for aufgabe in aufgaben]
                    
                    right_column_content = [
                        Paragraph(unternehmen, self.custom_styles['Position']),
                        Paragraph(position, self.custom_styles['Company'])
                    ]
                    right_column_content.extend(aufgaben_formatted)
                    
                    zeitraum_formatted = zeitraum
                    
                    data = [[Paragraph(zeitraum_formatted, self.custom_styles['Period']), right_column_content[0]]]
                    data.append([Paragraph('', self.custom_styles['Normal']), right_column_content[1]])
                    
                    if len(right_column_content) > 2:
                        data.append([Paragraph('', self.custom_styles['Normal']), Spacer(1, 5)])
                    
                    for j in range(2, len(right_column_content)):
                        data.append([Paragraph('', self.custom_styles['Normal']), right_column_content[j]])
                    
                    col_widths = [A4[0] * 0.15, A4[0] * 0.65]
                    
                    table = Table(data, colWidths=col_widths)
                    table.setStyle(TableStyle([
                        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                        ('LEFTPADDING', (1, 0), (1, -1), 0.5*cm),
                        ('BOTTOMPADDING', (0, 0), (0, 0), 0),
                        ('TOPPADDING', (1, 1), (1, 1), 1),
                    ]))
                    
                    gruppen_elemente.extend([table, Spacer(1, 0.5*cm)])
                
                try:
                    if len(gruppe) <= 3:
                        elements.append(KeepTogether(gruppen_elemente))
                    else:
                        elements.extend(gruppen_elemente)
                except Exception as e:
                    print(f"Fehler beim Verarbeiten einer Berufsgruppe: {str(e)}")
                    elements.extend(gruppen_elemente)
        
        elements.append(Spacer(1, 0.5*cm))
        elements.append(HRFlowable(width="100%", thickness=1, lineCap='round', color=colors.lightgrey))
        elements.append(Spacer(1, 0.5*cm))
        
        ausbildungen = profile_data.get('ausbildung', [])
        if ausbildungen:
            ausbildung_elements = []
            ausbildung_elements.append(Paragraph("Ausbildung", self.custom_styles['Heading2']))
            
            for i, ausbildung in enumerate(ausbildungen):
                try:
                    zeitraum = ausbildung.get('zeitraum', '')
                    if "bis jetzt" in zeitraum.lower() or "bis heute" in zeitraum.lower():
                        zeitraum = zeitraum.replace("bis jetzt", "2025").replace("bis heute", "2025").replace("bis JETZT", "2025")
                    institution = ausbildung.get('institution', '')
                    abschluss = ausbildung.get('abschluss', '')
                    schwerpunkte = ausbildung.get('schwerpunkte', '')
                    
                    right_column_content = []
                    right_column_content.append(Paragraph(institution, self.custom_styles['Position']))
                    
                    if abschluss:
                        right_column_content.append(Paragraph(f"{abschluss}", self.custom_styles['Company']))
                    
                    aufgaben = []
                    if schwerpunkte:
                        schwerpunkte_liste = schwerpunkte.split(", ")
                        for schwerpunkt in schwerpunkte_liste:
                            aufgaben.append(f"• {schwerpunkt}")
                    
                    if 'aufgaben' in ausbildung and isinstance(ausbildung['aufgaben'], list):
                        for aufgabe in ausbildung['aufgaben']:
                            aufgaben.append(f"• {aufgabe}")
                    
                    for aufgabe in aufgaben:
                        right_column_content.append(Paragraph(aufgabe, self.custom_styles['Normal']))
                    
                    zeitraum_formatted = zeitraum
                    
                    data = [[Paragraph(zeitraum_formatted, self.custom_styles['Period']), right_column_content[0]]]
                    
                    if len(right_column_content) > 1:
                        data.append([Paragraph('', self.custom_styles['Normal']), right_column_content[1]])
                    
                    for j in range(2, len(right_column_content)):
                        data.append([Paragraph('', self.custom_styles['Normal']), right_column_content[j]])
                    
                    col_widths = [A4[0] * 0.15, A4[0] * 0.65]
                    
                    table = Table(data, colWidths=col_widths)
                    table.setStyle(TableStyle([
                        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                        ('LEFTPADDING', (1, 0), (1, -1), 0.5*cm),
                        ('BOTTOMPADDING', (0, 0), (0, 0), 0),
                        ('TOPPADDING', (1, 1), (1, 1), 1),
                    ]))
                    
                    if i == 0:
                        ausbildung_elements.extend([table, Spacer(1, 0.1*cm)])
                    else:
                        elements.extend(ausbildung_elements)
                        ausbildung_elements = []
                        entry_elements = [table, Spacer(1, 0.5*cm)]
                        elements.append(KeepTogether(entry_elements))
                        
                except Exception as e:
                    print(f"Fehler bei der Verarbeitung einer Ausbildung: {str(e)}")
                    fallback_element = Paragraph(f"{zeitraum} - {institution}", self.custom_styles['Normal'])
                    if i == 0:
                        ausbildung_elements.extend([fallback_element, Spacer(1, 0.1*cm)])
                    else:
                        elements.extend(ausbildung_elements)
                        ausbildung_elements = []
                        elements.append(fallback_element)
                        elements.append(Spacer(1, 0.1*cm))
            
            if ausbildung_elements:
                elements.append(KeepTogether(ausbildung_elements))
        
        weiterbildungen = profile_data.get('weiterbildungen', [])
        if weiterbildungen:
            elements.append(Spacer(1, 0.5*cm))
            elements.append(HRFlowable(width="100%", thickness=0.5, lineCap='round', color=colors.lightgrey, spaceBefore=0.3*cm, spaceAfter=0.3*cm))
            elements.append(Spacer(1, 0.3*cm))
            
            weiterbildung_elements = []
            weiterbildung_elements.append(Paragraph("Fort- und Weiterbildungen", self.custom_styles['Heading2']))
            
            for i, weiterbildung in enumerate(weiterbildungen):
                try:
                    zeitraum = weiterbildung.get('zeitraum', '')
                    if "bis jetzt" in zeitraum.lower() or "bis heute" in zeitraum.lower():
                        zeitraum = zeitraum.replace("bis jetzt", "2025").replace("bis heute", "2025").replace("bis JETZT", "2025")
                    bezeichnung = weiterbildung.get('bezeichnung', '')
                    abschluss = weiterbildung.get('abschluss', '')
                    
                    right_column_content = []
                    
                    bezeichnung_clean = bezeichnung.replace("zur ", "").replace("zum ", "")
                    right_column_content.append(Paragraph(f"Fortbildung zum {bezeichnung_clean}", self.custom_styles['Position']))
                    
                    if abschluss and abschluss not in bezeichnung:
                        right_column_content.append(Paragraph(f"{abschluss}", self.custom_styles['Company']))
                    
                    aufgaben = []
                    
                    if 'inhalte' in weiterbildung and isinstance(weiterbildung['inhalte'], list):
                        for inhalt in weiterbildung['inhalte']:
                            aufgaben.append(f"• {inhalt}")
                    
                    if 'aufgaben' in weiterbildung and isinstance(weiterbildung['aufgaben'], list):
                        for aufgabe in weiterbildung['aufgaben']:
                            aufgaben.append(f"• {aufgabe}")
                    
                    for aufgabe in aufgaben:
                        right_column_content.append(Paragraph(aufgabe, self.custom_styles['Normal']))
                    
                    zeitraum_formatted = zeitraum
                    
                    data = [[Paragraph(zeitraum_formatted, self.custom_styles['Period']), right_column_content[0]]]
                    
                    if len(right_column_content) > 1:
                        data.append([Paragraph('', self.custom_styles['Normal']), right_column_content[1]])
                    
                    for j in range(2, len(right_column_content)):
                        data.append([Paragraph('', self.custom_styles['Normal']), right_column_content[j]])
                    
                    col_widths = [A4[0] * 0.15, A4[0] * 0.65]
                    
                    table = Table(data, colWidths=col_widths)
                    table.setStyle(TableStyle([
                        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                        ('LEFTPADDING', (1, 0), (1, -1), 0.5*cm),
                        ('BOTTOMPADDING', (0, 0), (0, 0), 0),
                        ('TOPPADDING', (1, 1), (1, 1), 1),
                    ]))
                    
                    if i == 0:
                        weiterbildung_elements.extend([table, Spacer(1, 0.1*cm)])
                    else:
                        elements.extend(weiterbildung_elements)
                        weiterbildung_elements = []
                        entry_elements = [table, Spacer(1, 0.5*cm)]
                        elements.append(KeepTogether(entry_elements))
                        
                except Exception as e:
                    print(f"Fehler bei der Verarbeitung einer Weiterbildung: {str(e)}")
                    fallback_element = Paragraph(f"{zeitraum} - {bezeichnung}", self.custom_styles['Normal'])
                    if i == 0:
                        weiterbildung_elements.extend([fallback_element, Spacer(1, 0.1*cm)])
                    else:
                        elements.extend(weiterbildung_elements)
                        weiterbildung_elements = []
                        elements.append(fallback_element)
                        elements.append(Spacer(1, 0.1*cm))
            
            if weiterbildung_elements:
                elements.append(KeepTogether(weiterbildung_elements))

    def _create_classic_docx_content(self, doc, profile_data):
        """Erstellt den Content für das Classic Template im DOCX Format"""
        
        personal_data = profile_data.get("persönliche_daten", {})
        
        name = personal_data.get("name", "")
        doc.add_paragraph(name).bold = True
        
        table = doc.add_table(rows=0, cols=2)
        table.autofit = True
        
        if personal_data.get("wohnort"):
            row_cells = table.add_row().cells
            row_cells[0].text = "Wohnort:"
            row_cells[0].paragraphs[0].runs[0].bold = True
            row_cells[1].text = personal_data.get("wohnort", "")
        
        if personal_data.get("jahrgang"):
            row_cells = table.add_row().cells
            row_cells[0].text = "Jahrgang:"
            row_cells[0].paragraphs[0].runs[0].bold = True
            row_cells[1].text = personal_data.get("jahrgang", "")
        
        if personal_data.get("führerschein"):
            row_cells = table.add_row().cells
            row_cells[0].text = "Führerschein:"
            row_cells[0].paragraphs[0].runs[0].bold = True
            row_cells[1].text = personal_data.get("führerschein", "")
        
        kontakt = personal_data.get('kontakt', {})
        ansprechpartner = kontakt.get('ansprechpartner', '')
        
        if ansprechpartner and ansprechpartner != "Kein Ansprechpartner":
            doc.add_paragraph().add_run().add_break()
            ansprechpartner_heading = doc.add_paragraph("IHR ANSPRECHPARTNER")
            ansprechpartner_heading.runs[0].bold = True
            ansprechpartner_heading.runs[0].font.size = Pt(9)
            
            doc.add_paragraph(ansprechpartner)
            
            if ansprechpartner == "Salim Alizai":
                telefon = kontakt.get('telefon', '')
            else:
                telefon = "02161 62126-00"
                
            if telefon:
                doc.add_paragraph(f"Tel.: {telefon}")
            
            email = kontakt.get('email', '')
            if email:
                doc.add_paragraph(f"E-Mail: {email}")
        
        doc.add_paragraph().add_run().add_break()
        heading_style = doc.styles.add_style('Heading2', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
        heading_style.font.size = Pt(12)
        heading_style.font.bold = True
        
        doc.add_paragraph("Beruflicher Werdegang", style='Heading2')
        
        for experience in profile_data.get("berufserfahrung", []):
            exp_table = doc.add_table(rows=0, cols=2)
            exp_table.autofit = True
            
            row_cells = exp_table.add_row().cells
            row_cells[0].text = experience.get("zeitraum", "")
            row_cells[0].paragraphs[0].runs[0].bold = True
            
            row_cells[1].text = experience.get("position", "")
            row_cells[1].paragraphs[0].runs[0].bold = True
            row_cells[1].add_paragraph(experience.get("unternehmen", experience.get("firma", ""))).italic = False
            
            if experience.get("beschreibung"):
                row_cells = exp_table.add_row().cells
                row_cells[0].merge(row_cells[1])
                row_cells[0].text = experience.get("beschreibung", "")
            
            doc.add_paragraph()
        
        doc.add_paragraph("Ausbildung", style='Heading2')
        
        for education in profile_data.get("ausbildung", []):
            edu_table = doc.add_table(rows=0, cols=2)
            edu_table.autofit = True
            
            row_cells = edu_table.add_row().cells
            row_cells[0].text = education.get("zeitraum", "")
            row_cells[0].paragraphs[0].runs[0].bold = True
            
            row_cells[1].text = education.get("institution", "")
            row_cells[1].paragraphs[0].runs[0].bold = True
            row_cells[1].add_paragraph(education.get("abschluss", "")).italic = False
            
            schwerpunkte = education.get("schwerpunkte", "")
            if schwerpunkte:
                schwerpunkte_liste = schwerpunkte.split(", ")
                for schwerpunkt in schwerpunkte_liste:
                    row_cells[1].add_paragraph(f"• {schwerpunkt}")
            
            if 'aufgaben' in education and isinstance(education['aufgaben'], list):
                for aufgabe in education['aufgaben']:
                    row_cells[1].add_paragraph(f"• {aufgabe}")
            
            doc.add_paragraph()
        
        doc.add_paragraph()
        doc.add_paragraph("Weiterbildungen", style='Heading2')
        
        if profile_data.get("weiterbildungen"):
            for training in profile_data.get("weiterbildungen", []):
                train_table = doc.add_table(rows=0, cols=2)
                train_table.autofit = True
                
                row_cells = train_table.add_row().cells
                row_cells[0].text = training.get("zeitraum", "")
                row_cells[0].paragraphs[0].runs[0].bold = True
                
                row_cells[1].text = training.get("bezeichnung", "")
                row_cells[1].paragraphs[0].runs[0].bold = True
                if training.get("abschluss"):
                    row_cells[1].add_paragraph(training.get("abschluss", "")).italic = False
                
                if 'inhalte' in training and isinstance(training['inhalte'], list):
                    for inhalt in training['inhalte']:
                        row_cells[1].add_paragraph(f"• {inhalt}")
                
                if 'aufgaben' in training and isinstance(training['aufgaben'], list):
                    for aufgabe in training['aufgaben']:
                        row_cells[1].add_paragraph(f"• {aufgabe}")
                
                doc.add_paragraph()
        
        doc.add_paragraph().add_run().add_break()
        footer_style = doc.styles.add_style('FooterStyle', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
        footer_style.font.size = Pt(8)
        footer_style.font.color.rgb = RGBColor(128, 128, 128)
        
        footer_text = self._get_dynamic_footer_text()
        doc.add_paragraph(footer_text, style='FooterStyle')

    def _create_custom_styles(self):
        """Erstellt benutzerdefinierte Stile für das Classic Template"""
        custom_styles = {}
        
        custom_styles['GaldoraLogo'] = ParagraphStyle(
            'GaldoraLogo', parent=self.styles['Normal'], fontSize=36, fontName='Helvetica-Bold',
            textColor=colors.black, spaceAfter=0.1*cm
        )
        custom_styles['Tagline'] = ParagraphStyle(
            'Tagline', parent=self.styles['Italic'], fontSize=10, fontName='Helvetica-Oblique',
            textColor=colors.black, spaceAfter=1.5*cm
        )
        custom_styles['ProfilTitle'] = ParagraphStyle(
            'ProfilTitle', parent=self.styles['Heading1'], fontSize=16, fontName='Helvetica',
            textColor=colors.HexColor('#1973B8'), spaceBefore=0.2*cm, spaceAfter=0.3*cm
        )
        custom_styles['Name'] = ParagraphStyle(
            'Name', parent=self.styles['Normal'], fontSize=13, fontName='Helvetica',
            spaceBefore=0.1*cm, spaceAfter=0.6*cm
        )
        custom_styles['Heading2'] = ParagraphStyle(
            'Heading2', parent=self.styles['Heading2'], fontSize=11, fontName='Helvetica-Bold',
            textColor=colors.black, spaceBefore=0.5*cm, spaceAfter=0.2*cm
        )
        custom_styles['ContactHeader'] = ParagraphStyle(
            'ContactHeader', parent=self.styles['Normal'], fontSize=8, fontName='Helvetica-Bold',
            spaceBefore=0.1*cm, spaceAfter=0.05*cm
        )
        custom_styles['ContactData'] = ParagraphStyle(
            'ContactData', parent=self.styles['Normal'], fontSize=8, textColor=colors.grey,
            spaceAfter=0.05*cm, leftIndent=0*cm
        )
        custom_styles['Normal'] = ParagraphStyle(
            'Normal', parent=self.styles['Normal'], fontSize=9, spaceAfter=0.01*cm
        )
        custom_styles['LabelInline'] = ParagraphStyle(
            'LabelInline', parent=self.styles['Normal'], fontSize=9, fontName='Helvetica-Bold',
            spaceAfter=0.05*cm
        )
        custom_styles['Company'] = ParagraphStyle(
            'Company', parent=self.styles['Normal'], fontSize=9, fontName='Helvetica',
            spaceAfter=0.01*cm
        )
        custom_styles['Position'] = ParagraphStyle(
            'Position', parent=self.styles['Normal'], fontSize=9, fontName='Helvetica-Bold',
            spaceAfter=0.01*cm
        )
        custom_styles['Period'] = ParagraphStyle(
            'Period', parent=self.styles['Normal'], fontSize=9, fontName='Helvetica-Bold',
            spaceAfter=0.05*cm
        )
        custom_styles['Footer'] = ParagraphStyle(
            'Footer', parent=self.styles['Normal'], fontSize=8, fontName='Helvetica',
            textColor=colors.grey, alignment=1
        )
        
        return custom_styles

    def _get_company_logo_path(self):
        """Gibt den Pfad zum Unternehmenslogo zurück, verlässt sich auf die utils."""
        try:
            logo_path = get_company_logo_path(self.selected_company)
            if logo_path and os.path.exists(logo_path):
                return logo_path
            
            company_config = self.company_config
            logo_filename = company_config.get('logo_filename', 'galdoralogo.png')
            
            fallback_path = get_image_path(logo_filename, use_static=True)
            if fallback_path and os.path.exists(fallback_path):
                return fallback_path
            
            print(f"Warnung: Logo nicht gefunden für Unternehmen: {self.selected_company}")
            return None
            
        except Exception as e:
            print(f"Fehler beim Laden des Logos: {str(e)}")
            return None

    def _get_dynamic_footer_text(self):
        """Gibt den dynamischen Footer-Text basierend auf dem ausgewählten Unternehmen zurück"""
        if self.selected_company == "bejob":
            return "bejob – einfach besser bewerben | Volksgartenstr. 85–89, 41065 Mönchengladbach | 02161 / 65326-40 | info@bejob.de | www.bejob.de"
        else:
            return "GALDORA Personalmanagement GmbH Co.KG\nVolksgartenstr. 85-89, 41065 Mönchengladbach\nE-Mail: info@galdora.de / Web: www.galdora.de"

def check_missing_profile_data(profile_data):
    """Überprüft, ob wichtige Felder in den Profildaten fehlen"""
    missing_fields = []
    
    personal_data = profile_data.get('persönliche_daten', {})
    if not personal_data.get('name'):
        missing_fields.append('Name')
    if not personal_data.get('wohnort'):
        missing_fields.append('Wohnort')
    if not personal_data.get('jahrgang'):
        missing_fields.append('Jahrgang/Geburtsdatum')
    
    if not profile_data.get('berufserfahrung', []):
        missing_fields.append('Berufserfahrung')
    
    if not profile_data.get('ausbildung', []):
        missing_fields.append('Ausbildung')
    
    return missing_fields
